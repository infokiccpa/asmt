import json
import pandas as pd
from docx import Document
import io

class QuizExporter:
    @staticmethod
    def to_json(questions: list) -> str:
        return json.dumps(questions, indent=4)
        
    @staticmethod
    def to_csv(questions: list) -> str:
        if not questions:
            return ""
        df = pd.DataFrame(questions)
        return df.to_csv(index=False)
        
    @staticmethod
    def to_docx_bytes(questions: list) -> bytes:
        doc = Document()
        doc.add_heading('Smart Quiz Extractor - Generated Quiz', 0)
        
        for idx, q in enumerate(questions, 1):
            doc.add_heading(f"Q{idx}:", level=1)
            doc.add_paragraph(q.get("question_text", ""))
            
            options = q.get("options", [])[:4]  # Limit to exactly 4 options
            for opt_idx, opt in enumerate(options):
                letter = chr(65 + opt_idx)
                doc.add_paragraph(f"  {letter}. {opt}")
                
            doc.add_paragraph(f"Correct Answer: {q.get('correct_answer', 'N/A')}")
            doc.add_paragraph("")
            
        byte_stream = io.BytesIO()
        doc.save(byte_stream)
        return byte_stream.getvalue()
